import time
from concurrent.futures import ProcessPoolExecutor, ThreadPoolExecutor
from pathlib import Path
from uuid import uuid4, UUID

import pymupdf
from docx import Document as DocxDocument
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

from app.models.document import Document


def text_loader(doc_path: str) -> Document:
    """Reads the text file, creates a document with the text.
    metadata includes type of the doc and token-size
    """
    with open(doc_path) as fp:
        data = fp.read()
    did = uuid4()
    metadata = {
        "type": "txt",
        "token-size": len(data.split()),
    }
    doc = Document(id=did, source=doc_path, text=data, metadata=metadata)
    return doc


def pdf_loader(doc_path: str) -> Document:
    """
    reads pdf files
    maintains the page number and offset map in metadata
    """
    full_text = ""
    page_offset_map = {}
    offset = 0
    pages = 0
    with pymupdf.open(doc_path) as doc:
        for page_num, page in enumerate(doc, start=1):
            page_text = page.get_text()
            full_text += page_text
            page_offset_map[page_num] = offset
            offset += len(page_text)
            pages += 1
    did = uuid4()
    metadata = {
        "type": "pdf",
        "num-pages": pages,
        "page-offsets": page_offset_map,
        "token-size": len(full_text.split()),
    }
    return Document(id=did, metadata=metadata, text=full_text, source=doc_path)


def _iter_blocks(doc):
    """Iterates the blocks of docx to ensure the paras and tables are in right order"""
    parent = doc.element.body

    for child in parent.iterchildren():

        if isinstance(child, CT_P):
            yield Paragraph(child, doc)

        elif isinstance(child, CT_Tbl):
            yield Table(child, doc)


def docx_loader(doc_path: str) -> Document:
    """
    reads docx files. will read tables + paragraphs
    """
    docx = DocxDocument(doc_path)
    parts = []
    for block in _iter_blocks(docx):
        if isinstance(block, Paragraph):
            parts.append(block.text)

        elif isinstance(block, Table):
            for row in block.rows:
                row_text = " | ".join(cell.text for cell in row.cells)
                parts.append(row_text)

    full_text = "\n".join(parts)
    metadata = {"type": "docx", "token-size": len(full_text.split())}
    return Document(id=uuid4(), source=doc_path, text=full_text, metadata=metadata)


loader_map = {
    "txt": text_loader,
    "pdf": pdf_loader,
    "docx": docx_loader,
    "md": text_loader,
}


def _parallel_load(file_path: Path) -> Document | None:
    extension = file_path.suffix.lower().lstrip(".")
    loader = loader_map.get(extension)
    if loader is None:
        print(f"WARNING: Unsupported file type: {file_path.name}")
        return
    try:
        document = loader(str(file_path))
        print(f"Loaded: {file_path.name}")
        return document

    except Exception as exc:
        print(f"WARNING: Failed to load " f"{file_path.name}: {exc}")
        return


def load_docs(data_path: str) -> list[Document]:
    documents = []
    paths = []
    for file_path in Path(data_path).iterdir():
        if not file_path.is_file():
            continue
        paths.append(file_path)
    s = time.perf_counter()
    with ProcessPoolExecutor() as executor:
        documents = executor.map(_parallel_load, paths)
    documents = [doc for doc in documents if doc is not None]
    print("loading took", (time.perf_counter() - s), "seconds")
    return documents


def create_doc_id_map(docs: list[Document]) -> dict[UUID, Document]:
    return {doc.id: doc for doc in docs}